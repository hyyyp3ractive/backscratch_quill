from docx import Document
from docx.shared import Inches
import datetime
import os

def build_mfr_docx(data):
    doc = Document()
    today = datetime.date.today().strftime("%d %b %Y")

    doc.add_heading("MEMORANDUM FOR RECORD", level=1)
    doc.add_paragraph(f"{today}")

    doc.add_paragraph(f"SUBJECT: {data.get('subject', 'Subject Here')}")
    doc.add_paragraph("")

    doc.add_paragraph(data.get("body", "Insert body text here."))

    doc.add_paragraph("")
    doc.add_paragraph("Respectfully,")
    doc.add_paragraph(data.get("signature_name", "[SIGNATORY NAME]"))
    doc.add_paragraph(data.get("signature_title", "[TITLE]"))

    output_path = os.path.join("exports", "output_mfr.docx")
    doc.save(output_path)
    return output_path
